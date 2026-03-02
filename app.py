import streamlit as st
from anthropic import Anthropic

# Streamlit сам найдет ключ в Secrets по этому имени
client = Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

# Пример вызова модели
# message = client.messages.create( ... )

import base64
import re
import io
from datetime import datetime
from PIL import Image

# ─── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ТехСпек Генератор",
    page_icon="⚙️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─── CSS — светлый профессиональный интерфейс ──────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif !important;
    font-size: 15px;
    color: #1a1a2e;
}

.stApp {
    background-color: #f5f7fa !important;
}

section[data-testid="stSidebar"],
div[data-testid="stAppViewContainer"],
div[data-testid="stHeader"] {
    background-color: #f5f7fa !important;
}

.app-header {
    background: linear-gradient(135deg, #1a1a2e 0%, #16213e 60%, #0f3460 100%);
    border-radius: 14px;
    padding: 32px 40px;
    margin-bottom: 28px;
    color: white;
}
.app-header h1 {
    font-size: 2rem;
    font-weight: 700;
    margin: 0 0 6px 0;
    color: #ffffff;
}
.app-header p { font-size: 1rem; color: #a8b4c8; margin: 0; }

.card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 24px 28px;
    margin-bottom: 18px;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}
.card-title {
    font-size: 0.75rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1.2px;
    color: #4361ee;
    margin: 0 0 16px 0;
}

.file-chip {
    background: #f0f4ff;
    border: 1px solid #c7d2fe;
    border-radius: 8px;
    padding: 9px 14px;
    margin: 5px 0;
    font-size: 0.85rem;
    color: #3730a3;
    font-weight: 500;
}

.badge { display: inline-block; padding: 4px 12px; border-radius: 20px; font-size: 0.75rem; font-weight: 600; }
.badge-green  { background: #dcfce7; color: #166534; border: 1px solid #86efac; }
.badge-blue   { background: #dbeafe; color: #1e40af; border: 1px solid #93c5fd; }

.spec-result {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 28px 32px;
    font-size: 0.95rem;
    line-height: 1.85;
    color: #1e293b;
    white-space: pre-wrap;
    box-shadow: 0 1px 4px rgba(0,0,0,0.06);
}

div[data-testid="stButton"] > button {
    background: linear-gradient(135deg, #4361ee, #3a0ca3) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    font-size: 1rem !important;
    padding: 12px 28px !important;
    box-shadow: 0 4px 12px rgba(67,97,238,0.3) !important;
}
div[data-testid="stButton"] > button:hover {
    box-shadow: 0 6px 18px rgba(67,97,238,0.45) !important;
}

div[data-testid="stDownloadButton"] > button {
    background: #ffffff !important;
    color: #1e40af !important;
    border: 1.5px solid #3b82f6 !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    padding: 9px 14px !important;
    width: 100% !important;
}
div[data-testid="stDownloadButton"] > button:hover {
    background: #eff6ff !important;
}

.stTextInput > div > div > input {
    background: #ffffff !important;
    border: 1.5px solid #cbd5e1 !important;
    border-radius: 8px !important;
    color: #1e293b !important;
    font-size: 0.95rem !important;
    padding: 10px 14px !important;
}
.stTextInput label, .stSelectbox label, .stSlider label {
    color: #374151 !important;
    font-weight: 500 !important;
    font-size: 0.9rem !important;
}

.stSelectbox > div > div {
    background: #ffffff !important;
    border: 1.5px solid #cbd5e1 !important;
    border-radius: 8px !important;
    color: #1e293b !important;
}

div[data-testid="stExpander"] {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 10px !important;
}

.footer {
    text-align: center;
    color: #94a3b8;
    font-size: 0.8rem;
    margin-top: 40px;
    padding-top: 20px;
    border-top: 1px solid #e2e8f0;
}

/* Поле промта */
.stTextArea > div > div > textarea {
    background: #ffffff !important;
    border: 1.5px solid #cbd5e1 !important;
    border-radius: 8px !important;
    color: #1e293b !important;
    font-size: 0.88rem !important;
    font-family: 'Courier New', monospace !important;
    line-height: 1.6 !important;
}
.stTextArea > div > div > textarea:focus {
    border-color: #4361ee !important;
    box-shadow: 0 0 0 3px rgba(67,97,238,0.1) !important;
}
.stTextArea label {
    color: #374151 !important;
    font-weight: 500 !important;
    font-size: 0.9rem !important;
}

/* Статистика токенов */
.token-stats {
    display: grid;
    grid-template-columns: 1fr 1fr 1fr 1fr;
    gap: 10px;
    margin: 14px 0 4px 0;
}
.token-box {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 10px;
    padding: 12px 10px;
    text-align: center;
}
.token-box .t-label {
    font-size: 0.65rem;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 0.8px;
    color: #94a3b8;
    margin-bottom: 4px;
}
.token-box .t-value {
    font-size: 1.05rem;
    font-weight: 700;
    color: #1e293b;
}
.token-box .t-sub {
    font-size: 0.65rem;
    color: #94a3b8;
    margin-top: 2px;
}
.token-box.hl-blue  { border-color: #93c5fd; background: #eff6ff; }
.token-box.hl-blue  .t-value { color: #1d4ed8; }
.token-box.hl-green { border-color: #86efac; background: #f0fdf4; }
.token-box.hl-green .t-value { color: #166534; }
.token-box.hl-cost  { border-color: #fcd34d; background: #fffbeb; }
.token-box.hl-cost  .t-value { color: #b45309; }
.token-box.hl-calls { border-color: #c4b5fd; background: #f5f3ff; }
.token-box.hl-calls .t-value { color: #6d28d9; }

/* Прогресс-бар */
.prog-wrap { background:#f1f5f9; border-radius:999px; height:12px; overflow:hidden; margin:8px 0 4px 0; }
.prog-fill  { height:100%; border-radius:999px; background:linear-gradient(90deg,#4361ee,#7c3aed); transition:width 0.4s; }
.prog-labels { font-size:0.78rem; color:#64748b; display:flex; justify-content:space-between; }
</style>
""", unsafe_allow_html=True)

# ─── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-header">
    <h1>⚙️ ТехСпек Генератор</h1>
    <p>Автоматическая генерация технических спецификаций для отдела закупа · PDF / DOCX / HTML · EN / ZH → RU</p>
</div>
""", unsafe_allow_html=True)

# ─── API Key ───────────────────────────────────────────────────────────────────
api_key = None
try:
    api_key = st.secrets["ANTHROPIC_API_KEY"]
except Exception:
    pass

if not api_key:
    with st.expander("🔑 API Ключ Anthropic", expanded=True):
        api_key_input = st.text_input(
            "Введи Anthropic API ключ",
            type="password",
            placeholder="sk-ant-api03-...",
        )
        st.markdown(
            "Получить ключ: [console.anthropic.com](https://console.anthropic.com) "
            "· При деплое на Streamlit Cloud добавь в **Settings → Secrets** как `ANTHROPIC_API_KEY`"
        )
        if api_key_input:
            api_key = api_key_input

# ─── Константы ─────────────────────────────────────────────────────────────────
SUPPORTED_TYPES = ["pdf", "png", "jpg", "jpeg", "webp", "tiff", "tif", "bmp"]
MAX_PDF_PAGES   = 90


# Цена claude-opus-4-5 на 1000 токенов ($ per 1K)
PRICE_INPUT_PER_1K  = 0.015   # $15 / 1M input
PRICE_OUTPUT_PER_1K = 0.075   # $75 / 1M output

DEFAULT_SYSTEM_PROMPT = """Ты — эксперт-технолог по технической документации для отдела закупа.

Твоя задача — анализировать документы (технические паспорта, datasheet'ы, каталоги) на любом языке и формировать точную, структурированную техническую спецификацию на РУССКОМ языке.

Правила:
- Переводи технические термины точно; оригинальные обозначения оставляй в скобках
- Не придумывай данные, которых нет в документе; пиши «не указано»
- Сохраняй все числовые значения и единицы измерения без изменений
- Структуру документа строго соблюдай (разделы 1–8)"""

# ─── Утилиты ───────────────────────────────────────────────────────────────────

def get_pdf_page_count(pdf_bytes):
    try:
        from pypdf import PdfReader
        return len(PdfReader(io.BytesIO(pdf_bytes)).pages)
    except Exception:
        return None


def split_pdf_chunks(pdf_bytes, max_pages=MAX_PDF_PAGES):
    """Split PDF into chunks ≤ max_pages. Returns list of (bytes, label)."""
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(io.BytesIO(pdf_bytes))
    total  = len(reader.pages)
    if total <= max_pages:
        return [(pdf_bytes, f"стр. 1–{total}")]
    chunks = []
    for start in range(0, total, max_pages):
        end    = min(start + max_pages, total)
        writer = PdfWriter()
        for i in range(start, end):
            writer.add_page(reader.pages[i])
        buf = io.BytesIO()
        writer.write(buf)
        chunks.append((buf.getvalue(), f"стр. {start+1}–{end}"))
    return chunks


def encode_image_for_claude(file_bytes, filename):
    name = filename.lower()
    if name.endswith((".tiff", ".tif", ".bmp")):
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return "image/png", base64.standard_b64encode(buf.getvalue()).decode()
    mt = {".png":"image/png",".jpg":"image/jpeg",".jpeg":"image/jpeg",".webp":"image/webp"}
    ext = "." + name.rsplit(".", 1)[-1]
    return mt.get(ext, "image/png"), base64.standard_b64encode(file_bytes).decode()


def make_preview(file_bytes, filename):
    try:
        if filename.lower().endswith(".pdf"):
            return None
        img = Image.open(io.BytesIO(file_bytes))
        img.thumbnail((700, 900))
        return img
    except Exception:
        return None


def build_prompt(equipment_type, doc_source_hint, detail_level, filenames, chunk_label=None):
    detail_map = {
        "Краткая":     "краткую — только ключевые параметры",
        "Стандартная": "стандартную с основными техническими характеристиками",
        "Подробная":   "максимально подробную со всеми параметрами и примечаниями",
    }
    eq_str     = f"Тип оборудования: {equipment_type}." if equipment_type else ""
    src_str    = doc_source_hint or ", ".join(filenames)
    chunk_note = f"\n[Обрабатывается фрагмент: {chunk_label}]" if chunk_label else ""

    return f"""Ты эксперт по технической документации отдела закупа.
Проанализируй документы (технические паспорта, документация производителя — могут быть на английском или китайском, в том числе сканы) и создай техническую спецификацию на РУССКОМ языке.
{chunk_note}
{eq_str}
Исходный документ: {src_str}
Файлы: {", ".join(filenames)}
Детализация: {detail_map.get(detail_level, "стандартную")}

Верни ТОЛЬКО текст спецификации. Никаких пояснений вне структуры.

## ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ

### 1. ВВЕДЕНИЕ
[Назначение и описание оборудования, 2-4 предложения]

### 2. ИСТОЧНИК ДОКУМЕНТАЦИИ
Исходный документ: {src_str}
Файл(ы): {", ".join(filenames)}

### 3. ОБЩИЕ СВЕДЕНИЯ
- Наименование: [полное наименование]
- Модель / Артикул: [если есть]
- Производитель: [компания, страна]
- Год / версия: [если указано]

### 4. ТЕХНИЧЕСКИЕ ХАРАКТЕРИСТИКИ
[Все параметры с единицами измерения]

### 5. УСЛОВИЯ ЭКСПЛУАТАЦИИ
[Температура, давление, среда, IP и т.д.]

### 6. КОМПЛЕКТАЦИЯ
[Что входит в поставку]

### 7. МОНТАЖ И ОБСЛУЖИВАНИЕ
[Ключевые требования]

### 8. СЕРТИФИКАТЫ И СТАНДАРТЫ
[Сертификаты, ГОСТ, ISO]

ПРАВИЛА: Весь текст только на русском. Если данных нет — «не указано». Ничего не придумывай."""


def call_claude(api_key, content_blocks, prompt_text, system_prompt=None, token_tracker=None):
    client  = anthropic.Anthropic(api_key=api_key)
    content = content_blocks + [{"type": "text", "text": prompt_text}]
    kwargs  = dict(
        model    = "claude-opus-4-5",
        max_tokens = 4096,
        messages = [{"role": "user", "content": content}]
    )
    if system_prompt and system_prompt.strip():
        kwargs["system"] = system_prompt.strip()

    resp = client.messages.create(**kwargs)

    # Track token usage
    if token_tracker is not None:
        usage = resp.usage
        token_tracker["input"]  = token_tracker.get("input", 0)  + usage.input_tokens
        token_tracker["output"] = token_tracker.get("output", 0) + usage.output_tokens
        token_tracker["calls"]  = token_tracker.get("calls", 0)  + 1

    return resp.content[0].text


def generate_spec(api_key, uploaded_files, equipment_type, doc_source_hint,
                  detail_level, progress_cb=None, system_prompt=None, token_tracker=None):
    filenames    = [f.name for f in uploaded_files]
    image_blocks = []
    pdf_chunks   = []

    for f in uploaded_files:
        raw = f.getvalue()
        if f.name.lower().endswith(".pdf"):
            for chunk_bytes, label in split_pdf_chunks(raw):
                pdf_chunks.append((chunk_bytes, label, f.name))
        else:
            mt, b64 = encode_image_for_claude(raw, f.name)
            image_blocks.append({"type":"image","source":{"type":"base64","media_type":mt,"data":b64}})

    prompt = build_prompt(equipment_type, doc_source_hint, detail_level, filenames)

    # Нет PDF — один вызов
    if not pdf_chunks:
        return call_claude(api_key, image_blocks, prompt, system_prompt, token_tracker)

    # Один чанк — один вызов
    if len(pdf_chunks) == 1:
        cb, lbl, _ = pdf_chunks[0]
        b64 = base64.standard_b64encode(cb).decode()
        blocks = image_blocks + [{"type":"document","source":{"type":"base64","media_type":"application/pdf","data":b64}}]
        return call_claude(api_key, blocks, prompt, system_prompt, token_tracker)

    # Несколько чанков — обрабатываем по очереди, потом объединяем
    partials = []
    total    = len(pdf_chunks)
    for idx, (cb, lbl, orig) in enumerate(pdf_chunks):
        if progress_cb:
            progress_cb(idx, total + 1, lbl)
        b64 = base64.standard_b64encode(cb).decode()
        blocks = image_blocks + [{"type":"document","source":{"type":"base64","media_type":"application/pdf","data":b64}}]
        chunk_prompt = build_prompt(equipment_type, doc_source_hint, detail_level,
                                    filenames, chunk_label=f"{orig} [{lbl}]")
        partials.append(f"=== Фрагмент {lbl} ===\n{call_claude(api_key, blocks, chunk_prompt, system_prompt, token_tracker)}")

    if progress_cb:
        progress_cb(total, total + 1, "объединяю фрагменты...")

    merge_prompt = (
        f"У тебя {len(partials)} частичных технических спецификации из разных фрагментов одного документа. "
        "Объедини их в ОДНУ итоговую спецификацию на русском языке. "
        "Убери дубликаты, сохрани все уникальные данные, структуру сохрани (разделы 1-8).\n\n"
        + "\n---\n".join(partials)
        + "\n\nВерни ТОЛЬКО итоговую спецификацию, без пояснений."
    )
    client = anthropic.Anthropic(api_key=api_key)
    resp   = client.messages.create(
        model="claude-opus-4-5", max_tokens=4096,
        messages=[{"role":"user","content":merge_prompt}]
    )
    if token_tracker is not None:
        token_tracker["input"]  = token_tracker.get("input", 0)  + resp.usage.input_tokens
        token_tracker["output"] = token_tracker.get("output", 0) + resp.usage.output_tokens
        token_tracker["calls"]  = token_tracker.get("calls", 0)  + 1
    return resp.content[0].text


def clean_filename(text, fallback="ТехСпек"):
    base  = text or fallback
    clean = re.sub(r'[^\w\s-]', '', base, flags=re.UNICODE)
    clean = re.sub(r'\s+', '_', clean.strip())[:50]
    return f"ТехСпек_{clean}_{datetime.now().strftime('%Y%m%d')}"


# ─── DOCX ГОСТ ─────────────────────────────────────────────────────────────────
def make_docx(spec_text, filename_base, doc_source_hint, filenames):
    from docx import Document as D
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = D()
    for s in doc.sections:
        s.page_width=Cm(21); s.page_height=Cm(29.7)
        s.left_margin=Cm(3); s.right_margin=Cm(1.5)
        s.top_margin=Cm(2);  s.bottom_margin=Cm(2)

    def sf(run, sz=12, bold=False, rgb=None):
        run.font.name="Times New Roman"; run.font.size=Pt(sz); run.font.bold=bold
        if rgb: run.font.color.rgb=RGBColor(*rgb)

    def sep():
        p=doc.add_paragraph()
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single')
        b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'1'); b.set(qn('w:color'),'003366')
        pBdr.append(b); pPr.append(pBdr); p.paragraph_format.space_after=Pt(6)

    src=doc_source_hint or ", ".join(filenames)
    for txt,sz,bold,rgb,align in [
        ("ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ",16,True,(0,0,102),WD_ALIGN_PARAGRAPH.CENTER),
        (f"на оборудование · {src}",11,False,(80,80,80),WD_ALIGN_PARAGRAPH.CENTER),
        (f"Дата: {datetime.now().strftime('%d.%m.%Y')}",10,False,(120,120,120),WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        p=doc.add_paragraph(); p.alignment=align
        sf(p.add_run(txt),sz,bold,rgb)
    sep()

    for line in spec_text.split('\n'):
        line=line.rstrip()
        if line.startswith('## ') or line.startswith('# '):
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(line.lstrip('#').strip().upper()); sf(r,13,True,(0,0,102))
            p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
        elif line.startswith('### '):
            p=doc.add_paragraph(); r=p.add_run(line.lstrip('#').strip()); sf(r,12,True,(0,51,102))
            p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
        elif line.startswith('- ') or line.startswith('* '):
            p=doc.add_paragraph(style='List Bullet'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            sf(p.add_run(line[2:].strip()),12); p.paragraph_format.space_after=Pt(2)
        elif line.startswith('---'):
            sep()
        elif line.strip()=='':
            doc.add_paragraph().paragraph_format.space_after=Pt(3)
        elif ':' in line and len(line)<200:
            parts=line.split(':',1); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            sf(p.add_run(parts[0].strip()+': '),12,True); sf(p.add_run(parts[1].strip()),12)
            p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=Pt(18)
        else:
            p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent=Cm(1.25); sf(p.add_run(line),12)
            p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=Pt(18)

    sep()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sf(p.add_run("Документ сформирован автоматически · ТехСпек Генератор"),9,rgb=(150,150,150))
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()


# ─── PDF ReportLab ─────────────────────────────────────────────────────────────
def make_pdf(spec_text, filename_base, doc_source_hint, filenames, style_num=1):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

    PAL={1:{"h1":colors.HexColor("#000066"),"h2":colors.HexColor("#003366"),"body":colors.black,"line":colors.HexColor("#1f6feb")},
         2:{"h1":colors.HexColor("#1a4731"),"h2":colors.HexColor("#2d5016"),"body":colors.HexColor("#1a1a1a"),"line":colors.HexColor("#238636")},
         3:{"h1":colors.black,"h2":colors.HexColor("#222"),"body":colors.black,"line":colors.HexColor("#555")}}
    FN={1:("Times-Roman","Times-Bold"),2:("Helvetica","Helvetica-Bold"),3:("Courier","Courier-Bold")}
    pal=PAL[style_num]; fn,fb=FN[style_num]
    src=doc_source_hint or ", ".join(filenames)

    buf=io.BytesIO()
    doc=SimpleDocTemplate(buf,pagesize=A4,leftMargin=3*cm,rightMargin=1.5*cm,topMargin=2*cm,bottomMargin=2*cm,title=filename_base)

    def PS(name,**kw): return ParagraphStyle(name,**kw)
    s_ttl =PS("t",fontName=fb,fontSize=16,leading=20,textColor=pal["h1"],alignment=TA_CENTER,spaceAfter=4)
    s_sub =PS("s",fontName=fn,fontSize=10,leading=14,textColor=colors.HexColor("#555"),alignment=TA_CENTER,spaceAfter=2)
    s_h1  =PS("h1",fontName=fb,fontSize=13,leading=17,textColor=pal["h1"],spaceBefore=12,spaceAfter=5)
    s_h2  =PS("h2",fontName=fb,fontSize=11,leading=15,textColor=pal["h2"],spaceBefore=8,spaceAfter=4)
    s_body=PS("b",fontName=fn,fontSize=11,leading=16,textColor=pal["body"],alignment=TA_JUSTIFY,firstLineIndent=20,spaceAfter=3)
    s_kv  =PS("kv",fontName=fn,fontSize=11,leading=15,textColor=pal["body"],spaceAfter=2)
    s_li  =PS("li",fontName=fn,fontSize=11,leading=15,textColor=pal["body"],leftIndent=16,spaceAfter=2)
    s_foot=PS("f",fontName=fn,fontSize=8,leading=11,textColor=colors.HexColor("#999"),alignment=TA_CENTER)

    story=[Paragraph("ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ",s_ttl),
           Paragraph(f"на оборудование · {src}",s_sub),
           Paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}",s_sub),
           HRFlowable(width="100%",thickness=1.5,color=pal["line"],spaceAfter=8)]

    for line in spec_text.split('\n'):
        line=line.rstrip()
        if   line.startswith('## ') or line.startswith('# '): story.append(Paragraph(line.lstrip('#').strip().upper(),s_h1))
        elif line.startswith('### '): story.append(Paragraph(line.lstrip('#').strip(),s_h2))
        elif line.startswith('- ') or line.startswith('* '): story.append(Paragraph(f"• {line[2:].strip()}",s_li))
        elif line.startswith('---'): story.append(HRFlowable(width="100%",thickness=0.5,color=pal["line"],spaceAfter=4))
        elif line.strip()=='': story.append(Spacer(1,4))
        elif ':' in line and len(line)<200:
            p=line.split(':',1); story.append(Paragraph(f"<b>{p[0].strip()}:</b> {p[1].strip()}",s_kv))
        else: story.append(Paragraph(line,s_body))

    story+=[Spacer(1,14),HRFlowable(width="100%",thickness=0.5,color=colors.HexColor("#ccc"),spaceAfter=4),
            Paragraph("Документ сформирован автоматически · ТехСпек Генератор",s_foot)]
    doc.build(story); buf.seek(0)
    return buf.getvalue()


def make_html(spec_text, filename_base, doc_source_hint, filenames, style_num=1):
    STYLES={
        1:"body{font-family:'Times New Roman',serif;max-width:820px;margin:40px auto;padding:0 48px;color:#111;font-size:14px;line-height:1.8;}h2{color:#000066;font-size:17px;text-align:center;text-transform:uppercase;border-bottom:2px solid #000066;padding-bottom:8px;}h3{color:#003366;font-size:14px;}hr{border:none;border-top:1px solid #ccc;margin:14px 0;}li{margin:3px 0;}p{text-align:justify;text-indent:28px;margin:4px 0;}",
        2:"body{font-family:Arial,sans-serif;max-width:820px;margin:40px auto;padding:0 48px;color:#1a1a1a;font-size:13.5px;line-height:1.75;}h2{color:#1a4731;font-size:16px;text-align:center;border-bottom:2px solid #238636;padding-bottom:8px;}h3{color:#2d5016;font-size:13px;}hr{border:none;border-top:1px solid #a8c795;margin:12px 0;}li{margin:3px 0;}p{text-align:justify;text-indent:24px;margin:4px 0;}",
        3:"body{font-family:'Courier New',monospace;max-width:820px;margin:40px auto;padding:0 48px;color:#111;font-size:12.5px;line-height:1.65;}h2{font-size:15px;text-transform:uppercase;border:2px solid #111;padding:8px;text-align:center;}h3{font-size:12.5px;text-transform:uppercase;border-left:4px solid #111;padding-left:8px;}hr{border:none;border-top:1px dashed #666;}li{margin:2px 0;}p{margin:3px 0;}",
    }
    src=doc_source_hint or ", ".join(filenames)
    css=STYLES.get(style_num,STYLES[1])
    parts=[]
    for line in spec_text.split('\n'):
        line=line.rstrip()
        if   line.startswith('## ') or line.startswith('# '): parts.append(f'<h2>{line.lstrip("#").strip()}</h2>')
        elif line.startswith('### '): parts.append(f'<h3>{line.lstrip("#").strip()}</h3>')
        elif line.startswith('- ') or line.startswith('* '): parts.append(f'<li>{line[2:]}</li>')
        elif line.startswith('---'): parts.append('<hr>')
        elif line.strip()=='': parts.append('<br>')
        elif ':' in line and len(line)<200:
            p=line.split(':',1); parts.append(f'<p><b>{p[0].strip()}:</b> {p[1].strip()}</p>')
        else: parts.append(f'<p>{line}</p>')

    return f"""<!DOCTYPE html>
<html lang="ru"><head><meta charset="UTF-8"><title>{filename_base}</title>
<style>{css}</style></head><body>
<div style="text-align:center;margin-bottom:20px;">
  <h2>Техническая спецификация</h2>
  <p style="color:#666;font-size:12px;">{src} · {datetime.now().strftime('%d.%m.%Y')}</p>
</div>
{''.join(parts)}
<hr><p style="text-align:center;font-size:10px;color:#aaa;margin-top:12px;">
Документ сформирован автоматически · ТехСпек Генератор</p>
</body></html>"""


# ─── Layout ────────────────────────────────────────────────────────────────────
col_left, col_right = st.columns([1, 1.5], gap="large")

with col_left:
    st.markdown('<div class="card"><div class="card-title">📂 Загрузка документов</div>', unsafe_allow_html=True)
    uploaded_files = st.file_uploader(
        "Перетащи файлы или нажми Browse",
        type=SUPPORTED_TYPES,
        accept_multiple_files=True,
        label_visibility="collapsed"
    )

    if uploaded_files:
        for f in uploaded_files:
            kb=len(f.getvalue())//1024; note=""
            if f.name.lower().endswith(".pdf"):
                pc=get_pdf_page_count(f.getvalue())
                if pc:
                    note=f" · {pc} стр."
                    if pc>MAX_PDF_PAGES:
                        n_chunks=-(-pc//MAX_PDF_PAGES)
                        note+=f" ⚠️ разобьётся на {n_chunks} части"
            st.markdown(f'<div class="file-chip">📄 {f.name} <span style="color:#6b7280">({kb} KB{note})</span></div>', unsafe_allow_html=True)
        st.markdown(f'<span class="badge badge-green">✓ {len(uploaded_files)} файл(ов)</span>', unsafe_allow_html=True)

        with st.expander("👁️ Предпросмотр загруженных файлов"):
            for f in uploaded_files:
                img=make_preview(f.getvalue(),f.name)
                if img:
                    st.markdown(f"**{f.name}**")
                    st.image(img,use_container_width=True)
                elif f.name.lower().endswith(".pdf"):
                    pc=get_pdf_page_count(f.getvalue())
                    st.info(f"📄 {f.name} — PDF, {pc} стр.")

    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="card"><div class="card-title">⚙️ Параметры генерации</div>', unsafe_allow_html=True)
    equipment_type  = st.text_input("Тип оборудования", placeholder="Насос, генератор, котёл...")
    doc_source_hint = st.text_input("Ссылка / наименование источника", placeholder="Технический паспорт Grundfos CM5-6, 2023")
    detail_level    = st.select_slider("Уровень детализации", options=["Краткая","Стандартная","Подробная"], value="Стандартная")
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Системный промт ──
    st.markdown('<div class="card"><div class="card-title">🧠 Системный промт (можно редактировать)</div>', unsafe_allow_html=True)
    st.markdown('<small style="color:#6b7280;">Это инструкция для AI — определяет как модель будет анализировать документы и писать спецификацию. Отредактируй под свои нужды.</small>', unsafe_allow_html=True)
    system_prompt = st.text_area(
        "Системный промт",
        value=DEFAULT_SYSTEM_PROMPT,
        height=200,
        label_visibility="collapsed",
        key="system_prompt_input"
    )
    col_reset, col_tokens = st.columns([1, 1])
    with col_reset:
        if st.button("↩️ Сбросить к умолчанию", use_container_width=True):
            st.session_state["system_prompt_input"] = DEFAULT_SYSTEM_PROMPT
            st.rerun()
    with col_tokens:
        st.markdown('<div style="font-size:0.78rem;color:#6b7280;padding:8px 0;">Модель: <b>claude-opus-4-5</b><br>Лимит: <b>200 000 токенов</b></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    can_go = bool(uploaded_files and api_key)
    gen_btn = st.button("⚡ Сгенерировать спецификацию", use_container_width=True, disabled=not can_go)
    if not api_key:       st.warning("⚠️ Введи API ключ выше")
    elif not uploaded_files: st.info("📂 Загрузи хотя бы один файл")

# ─── Результат ─────────────────────────────────────────────────────────────────
with col_right:
    if gen_btn and can_go:
        MAX_TOKENS = 200_000
        token_tracker = {"input": 0, "output": 0, "calls": 0}

        # Progress area
        st.markdown('<div class="card"><div class="card-title">⏳ Прогресс выполнения</div>', unsafe_allow_html=True)
        prog_bar_html = st.empty()
        prog_status   = st.empty()
        token_display = st.empty()
        st.markdown('</div>', unsafe_allow_html=True)

        def render_progress(pct, status_text, tracker):
            pct = min(pct, 100)
            used   = tracker.get("input", 0) + tracker.get("output", 0)
            cost_i = tracker.get("input",  0) / 1000 * PRICE_INPUT_PER_1K
            cost_o = tracker.get("output", 0) / 1000 * PRICE_OUTPUT_PER_1K
            cost   = cost_i + cost_o

            prog_bar_html.markdown(f"""
            <div class="prog-wrap">
                <div class="prog-fill" style="width:{pct}%"></div>
            </div>
            <div class="prog-labels">
                <span>{status_text}</span>
                <span style="font-weight:600;color:#4361ee">{pct}%</span>
            </div>
            """, unsafe_allow_html=True)

            token_display.markdown(f"""
            <div class="token-stats">
                <div class="token-box hl-blue">
                    <div class="t-label">📥 Input</div>
                    <div class="t-value">{tracker.get("input",0):,}</div>
                    <div class="t-sub">токенов</div>
                </div>
                <div class="token-box hl-green">
                    <div class="t-label">📤 Output</div>
                    <div class="t-value">{tracker.get("output",0):,}</div>
                    <div class="t-sub">токенов</div>
                </div>
                <div class="token-box hl-calls">
                    <div class="t-label">🔄 Запросов</div>
                    <div class="t-value">{tracker.get("calls",0)}</div>
                    <div class="t-sub">API calls</div>
                </div>
                <div class="token-box hl-cost">
                    <div class="t-label">💰 Стоимость</div>
                    <div class="t-value">${cost:.4f}</div>
                    <div class="t-sub">USD</div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        render_progress(2, "Подготовка файлов...", token_tracker)

        def pcb(done, total, label):
            pct = int(done / max(total, 1) * 95)
            render_progress(pct, f"Обрабатываю: {label} ({done}/{total})", token_tracker)

        try:
            spec = generate_spec(
                api_key, uploaded_files, equipment_type, doc_source_hint,
                detail_level, progress_cb=pcb,
                system_prompt=st.session_state.get("system_prompt_input", DEFAULT_SYSTEM_PROMPT),
                token_tracker=token_tracker
            )
            render_progress(100, "✅ Готово!", token_tracker)
            st.session_state.update({
                "spec": spec,
                "fn": clean_filename(equipment_type or doc_source_hint),
                "src": doc_source_hint,
                "fnames": [f.name for f in uploaded_files],
                "token_tracker": dict(token_tracker),
            })
            st.success("✅ Спецификация сгенерирована!")
        except Exception as e:
            render_progress(0, f"❌ Ошибка", token_tracker)
            st.error(f"❌ Ошибка: {e}")

    if "spec" in st.session_state:
        spec=st.session_state["spec"]; fn=st.session_state["fn"]
        src=st.session_state["src"];   fnames=st.session_state["fnames"]
        tracker=st.session_state.get("token_tracker", {})

        # Token stats (persistent after generation)
        if tracker:
            cost_i = tracker.get("input",  0) / 1000 * PRICE_INPUT_PER_1K
            cost_o = tracker.get("output", 0) / 1000 * PRICE_OUTPUT_PER_1K
            cost   = cost_i + cost_o
            used   = tracker.get("input", 0) + tracker.get("output", 0)
            pct    = min(int(used / 200_000 * 100), 100)
            st.markdown(f"""
            <div class="card">
                <div class="card-title">📊 Статистика последнего запроса</div>
                <div class="prog-wrap">
                    <div class="prog-fill" style="width:{pct}%"></div>
                </div>
                <div class="prog-labels">
                    <span>Использовано токенов</span>
                    <span style="font-weight:600;">{used:,} / 200 000 &nbsp;({pct}%)</span>
                </div>
                <div class="token-stats" style="margin-top:12px;">
                    <div class="token-box hl-blue">
                        <div class="t-label">📥 Input</div>
                        <div class="t-value">{tracker.get("input",0):,}</div>
                        <div class="t-sub">токенов</div>
                    </div>
                    <div class="token-box hl-green">
                        <div class="t-label">📤 Output</div>
                        <div class="t-value">{tracker.get("output",0):,}</div>
                        <div class="t-sub">токенов</div>
                    </div>
                    <div class="token-box hl-calls">
                        <div class="t-label">🔄 Запросов</div>
                        <div class="t-value">{tracker.get("calls",0)}</div>
                        <div class="t-sub">API calls</div>
                    </div>
                    <div class="token-box hl-cost">
                        <div class="t-label">💰 Стоимость</div>
                        <div class="t-value">${cost:.4f}</div>
                        <div class="t-sub">≈ {cost*500:.1f} тнг</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown('<div class="card"><div class="card-title">📄 Результат</div>', unsafe_allow_html=True)
        st.markdown(f'<span class="badge badge-blue">📁 {fn}</span><br><br>', unsafe_allow_html=True)
        st.markdown(f'<div class="spec-result">{spec}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="card"><div class="card-title">⬇️ Скачать документ</div>', unsafe_allow_html=True)

        st.markdown("**📝 Word (.docx) — стиль ГОСТ**")
        st.download_button("⬇️ Скачать DOCX (ГОСТ)", make_docx(spec,fn,src,fnames),
                           f"{fn}.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           use_container_width=True, key="docx")

        st.markdown("<br>**📑 PDF — 3 варианта оформления**", unsafe_allow_html=True)
        c1,c2,c3=st.columns(3)
        for col,(num,lbl) in zip([c1,c2,c3],[(1,"Корпоративный\n(Times)"),(2,"Официальный\n(Arial)"),(3,"Технический\n(Courier)")]):
            with col:
                st.download_button(f"⬇️ PDF #{num}", make_pdf(spec,fn,src,fnames,num), f"{fn}_v{num}.pdf","application/pdf",key=f"p{num}",use_container_width=True)
                st.caption(lbl)

        st.markdown("<br>**🌐 HTML — 3 варианта**", unsafe_allow_html=True)
        h1,h2,h3=st.columns(3)
        for col,(num,lbl) in zip([h1,h2,h3],[(1,"Корпоративный"),(2,"Официальный"),(3,"Технический")]):
            with col:
                st.download_button(f"⬇️ HTML #{num}", make_html(spec,fn,src,fnames,num).encode("utf-8"),f"{fn}_v{num}.html","text/html",key=f"h{num}",use_container_width=True)
                st.caption(lbl)

        st.markdown("<br>", unsafe_allow_html=True)
        st.download_button("📄 TXT",spec.encode("utf-8"),f"{fn}.txt","text/plain",use_container_width=True,key="txt")
        st.markdown('</div>', unsafe_allow_html=True)

    elif not gen_btn:
        st.markdown("""
        <div class="card" style="padding:48px 32px;">
            <div class="card-title">Как пользоваться</div>
            <div style="font-size:1rem;line-height:2.4;color:#374151;">
                1️⃣ &nbsp;Введи API ключ Anthropic<br>
                2️⃣ &nbsp;Загрузи файлы — PDF, фото, TIFF, сканы<br>
                3️⃣ &nbsp;Укажи тип оборудования и источник<br>
                4️⃣ &nbsp;Нажми <strong>«Сгенерировать спецификацию»</strong><br>
                5️⃣ &nbsp;Скачай в нужном формате
            </div>
            <br>
            <div style="font-size:0.85rem;color:#6b7280;background:#f8fafc;border-radius:8px;padding:14px 18px;border:1px solid #e2e8f0;line-height:2;">
                📎 Форматы входных файлов: PDF · PNG · JPG · WEBP · TIFF · BMP<br>
                🌐 Языки документов: English · 中文 → Русский<br>
                📤 Выходные форматы: DOCX (ГОСТ) · PDF ×3 · HTML ×3 · TXT<br>
                📄 PDF больше 100 страниц — разбивается автоматически
            </div>
        </div>
        """, unsafe_allow_html=True)

st.markdown('<div class="footer">ТехСпек Генератор · Claude API · для отдела закупа</div>', unsafe_allow_html=True)
